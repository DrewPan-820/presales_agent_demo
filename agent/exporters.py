"""文档导出工具：把架构师输出的 Markdown 方案转成 Word/纯文本。

只覆盖架构师真实会用到的 Markdown 子集：
  - # / ## / ### 标题
  - 普通段落
  - 无序列表（- / *）
  - 有序列表（1. 2.）
  - Markdown 表格（含分隔行）
  - 行内 **加粗**

不依赖 pandoc，纯 Python 实现，体积小。
"""
from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_ORDERED_PATTERN = re.compile(r"^\d+\.\s+")
_TABLE_SEP_PATTERN = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """把含 **bold** 的文本切分后写入 paragraph，保持加粗格式。"""
    cursor = 0
    for match in _BOLD_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _split_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip() for cell in raw.split("|")]


def markdown_to_docx(md_text: str, *, title: str | None = None) -> bytes:
    """把 Markdown 文本转成 .docx 字节流。"""
    doc = Document()

    # 默认正文字体（Word 中文显示更友好）
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # ---- Markdown 表格 ----
        if line.lstrip().startswith("|") and i + 1 < len(lines) and _TABLE_SEP_PATTERN.match(lines[i + 1]):
            header = _split_table_row(line)
            i += 2  # 跳过表头行 + 分隔行
            body_rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                body_rows.append(_split_table_row(lines[i]))
                i += 1

            ncols = len(header)
            table = doc.add_table(rows=1 + len(body_rows), cols=ncols)
            try:
                table.style = "Light Grid Accent 1"
            except KeyError:
                table.style = "Table Grid"

            for c, cell_text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True

            for r, row in enumerate(body_rows, start=1):
                for c in range(ncols):
                    cell = table.rows[r].cells[c]
                    cell.text = ""
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    text = row[c] if c < len(row) else ""
                    _add_runs_with_bold(cell.paragraphs[0], text)

            doc.add_paragraph()  # 表格后空行
            continue

        # ---- 标题 ----
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # ---- 列表 ----
        elif line.lstrip().startswith(("- ", "* ")):
            content = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, content)
        elif _ORDERED_PATTERN.match(line.lstrip()):
            content = _ORDERED_PATTERN.sub("", line.lstrip(), count=1)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, content)
        # ---- 普通段落 ----
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
